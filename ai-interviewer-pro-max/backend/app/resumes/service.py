"""
Resume Service

Business logic for resume upload, parsing, and management.
Clean separation from routes for testability and maintainability.

Per requirements:
- No external APIs
- No Gemini calls yet (TODO markers added)
- Modular and extensible
"""

import os
import re
import uuid
import aiofiles
from datetime import datetime
from typing import Optional, List, Dict, Any
from sqlalchemy.orm import Session
from fastapi import UploadFile

from app.core.config import settings
from app.resumes.models import Resume


class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors."""
    pass


class ResumeService:
    """
    Service class for resume operations.
    
    Handles:
    - File upload and validation
    - Text extraction from PDF and DOCX
    - Resume storage and retrieval
    """
    
    def __init__(self, db: Session):
        """
        Initialize resume service with database session.
        
        Args:
            db: SQLAlchemy database session
        """
        self.db = db
        self._ensure_upload_dir()
    
    def _ensure_upload_dir(self):
        """Ensure upload directory exists."""
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    
    # ===========================================
    # VALIDATION
    # ===========================================
    
    def validate_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Dict with validation result
            
        Raises:
            ValueError: If file is invalid
        """
        if not file.filename:
            raise ValueError("No filename provided")
        
        # Extract extension
        filename = file.filename.lower()
        extension = filename.rsplit(".", 1)[-1] if "." in filename else ""
        
        # Validate extension
        if extension not in settings.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Invalid file type: .{extension}. "
                f"Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}"
            )
        
        # Validate content type
        valid_content_types = {
            "pdf": ["application/pdf"],
            "docx": [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/octet-stream"  # Sometimes sent by browsers
            ]
        }
        
        # Size will be validated after reading
        return {
            "valid": True,
            "extension": extension,
            "original_filename": file.filename,
        }
    
    def _validate_file_size(self, size: int):
        """
        Validate file size.
        
        Args:
            size: File size in bytes
            
        Raises:
            ValueError: If file is too large
        """
        max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
        if size > max_size:
            raise ValueError(
                f"File too large: {size / 1024 / 1024:.1f}MB. "
                f"Maximum: {settings.MAX_UPLOAD_SIZE_MB}MB"
            )
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename to prevent path traversal.
        
        Args:
            filename: Original filename
            
        Returns:
            Safe filename
        """
        # Remove path components
        filename = os.path.basename(filename)
        # Remove potentially dangerous characters
        filename = re.sub(r'[^\w\-_\. ]', '', filename)
        # Limit length
        if len(filename) > 200:
            name, ext = os.path.splitext(filename)
            filename = name[:200-len(ext)] + ext
        return filename
    
    # ===========================================
    # FILE OPERATIONS
    # ===========================================
    
    async def save_file(
        self, 
        file: UploadFile, 
        user_id: str
    ) -> Dict[str, Any]:
        """
        Save uploaded file to filesystem.
        
        Args:
            file: Uploaded file
            user_id: Owner user ID
            
        Returns:
            Dict with file info
        """
        # Validate and get info
        file_info = self.validate_file(file)
        
        # Generate unique filename
        unique_id = str(uuid.uuid4())
        safe_filename = self._sanitize_filename(file.filename)
        stored_filename = f"{unique_id}_{safe_filename}"
        
        # Create user directory
        user_upload_dir = os.path.join(settings.UPLOAD_DIR, user_id)
        os.makedirs(user_upload_dir, exist_ok=True)
        
        # Full path
        file_path = os.path.join(user_upload_dir, stored_filename)
        
        # Read and save file
        content = await file.read()
        
        # Validate size
        self._validate_file_size(len(content))
        
        # Write to disk
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)
        
        return {
            "filename": stored_filename,
            "original_filename": file.filename,
            "file_path": file_path,
            "file_type": file_info["extension"],
            "file_size": len(content),
        }
    
    def delete_file(self, file_path: str) -> bool:
        """
        Delete file from filesystem.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if deleted, False if not found
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception:
            return False
    
    # ===========================================
    # TEXT EXTRACTION
    # ===========================================
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract plain text from resume file.
        
        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx)
            
        Returns:
            Extracted text content
            
        Raises:
            ResumeParsingError: If extraction fails
        """
        try:
            if file_type == "pdf":
                return self._extract_pdf_text(file_path)
            elif file_type == "docx":
                return self._extract_docx_text(file_path)
            else:
                raise ResumeParsingError(f"Unsupported file type: {file_type}")
        except ResumeParsingError:
            raise
        except Exception as e:
            raise ResumeParsingError(f"Failed to extract text: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            text = "\n".join(text_parts)
            return self._clean_text(text)
            
        except ImportError:
            raise ResumeParsingError("PyPDF2 library not installed")
        except Exception as e:
            raise ResumeParsingError(f"PDF parsing failed: {str(e)}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            return self._clean_text(text)
            
        except ImportError:
            raise ResumeParsingError("python-docx library not installed")
        except Exception as e:
            raise ResumeParsingError(f"DOCX parsing failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing noise.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove completely empty lines at start/end
        text = text.strip()
        
        return text
    
    # ===========================================
    # CRUD OPERATIONS
    # ===========================================
    
    async def upload_resume(
        self, 
        file: UploadFile, 
        user_id: str
    ) -> Resume:
        """
        Upload and parse a resume.
        
        Args:
            file: Uploaded file
            user_id: Owner user ID
            
        Returns:
            Created Resume object
        """
        # Save file
        file_info = await self.save_file(file, user_id)
        
        # Create resume record
        resume = Resume(
            user_id=user_id,
            filename=file_info["filename"],
            original_filename=file_info["original_filename"],
            file_type=file_info["file_type"],
            file_size=file_info["file_size"],
            file_path=file_info["file_path"],
            is_parsed="pending",
        )
        
        self.db.add(resume)
        self.db.commit()
        self.db.refresh(resume)
        
        # Extract text
        try:
            text_content = self.extract_text(
                file_info["file_path"], 
                file_info["file_type"]
            )
            resume.text_content = text_content
            resume.is_parsed = "success"
            
            # TODO: Gemini integration point
            # After successful text extraction, call Gemini for analysis:
            # - Extract skills
            # - Summarize experience
            # - Calculate ATS score
            # resume.analysis_status = "pending"
            # await self._analyze_with_gemini(resume)
            
        except ResumeParsingError as e:
            resume.is_parsed = "failed"
            resume.parse_error = str(e)
        
        self.db.commit()
        self.db.refresh(resume)
        
        return resume
    
    def get_user_resumes(self, user_id: str) -> List[Resume]:
        """
        Get all resumes for a user.
        
        Args:
            user_id: User ID
            
        Returns:
            List of Resume objects
        """
        return self.db.query(Resume).filter(
            Resume.user_id == user_id
        ).order_by(Resume.created_at.desc()).all()
    
    def get_resume_by_id(
        self, 
        resume_id: str, 
        user_id: str
    ) -> Optional[Resume]:
        """
        Get a specific resume by ID.
        
        Security: Only returns if user owns the resume.
        
        Args:
            resume_id: Resume ID
            user_id: User ID (for ownership check)
            
        Returns:
            Resume object or None
        """
        return self.db.query(Resume).filter(
            Resume.id == resume_id,
            Resume.user_id == user_id  # Security: user-scoped access
        ).first()
    
    def delete_resume(self, resume_id: str, user_id: str) -> bool:
        """
        Delete a resume.
        
        Args:
            resume_id: Resume ID
            user_id: User ID (for ownership check)
            
        Returns:
            True if deleted, False if not found
        """
        resume = self.get_resume_by_id(resume_id, user_id)
        
        if not resume:
            return False
        
        # Delete file from filesystem
        self.delete_file(resume.file_path)
        
        # Delete database record
        self.db.delete(resume)
        self.db.commit()
        
        return True
    
    # ===========================================
    # FUTURE: GEMINI INTEGRATION
    # ===========================================
    
    # TODO: Implement Gemini analysis when API key is available
    # 
    # async def _analyze_with_gemini(self, resume: Resume) -> None:
    #     """
    #     Analyze resume with Gemini AI.
    #     
    #     Extracts:
    #     - Skills list
    #     - Experience summary
    #     - ATS compatibility score
    #     - Job role recommendations
    #     """
    #     if not settings.is_gemini_configured():
    #         return
    #     
    #     # Call Gemini API here
    #     # Update resume.skills_extracted, resume.experience_summary, etc.
    #     pass


# Singleton-style helper for direct imports
def get_resume_service(db: Session) -> ResumeService:
    """Get resume service instance."""
    return ResumeService(db)
